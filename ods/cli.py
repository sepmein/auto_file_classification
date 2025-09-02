"""
命令行界面

提供基础的命令行功能
"""

import click
import logging
from pathlib import Path
from typing import Optional

from .core.config import Config
from .core.workflow import DocumentClassificationWorkflow
from .parsers.document_parser import DocumentParser


def setup_logging(log_level: str = "INFO") -> None:
    """设置日志"""
    logging.basicConfig(
        level=getattr(logging, log_level.upper()),
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    )


@click.group()
@click.option("--config", "-c", help="配置文件路径")
@click.option("--verbose", "-v", is_flag=True, help="详细输出")
@click.pass_context
def main(ctx, config: Optional[str], verbose: bool):
    """基于LLM和向量数据库的自动文档分类系统"""

    # 设置日志
    log_level = "DEBUG" if verbose else "INFO"
    setup_logging(log_level)

    # 加载配置
    ctx.ensure_object(dict)
    ctx.obj["config"] = Config(config)

    if verbose:
        click.echo(f"配置文件: {ctx.obj['config'].config_path}")


@main.command()
@click.pass_context
def init(ctx):
    """初始化系统"""
    config = ctx.obj["config"]

    click.echo("正在初始化系统...")

    # 创建必要的目录
    for directory in [config.system.temp_directory, Path(config.database.path).parent]:
        Path(directory).mkdir(parents=True, exist_ok=True)
        click.echo(f"创建目录: {directory}")

    # 保存配置
    config.save()
    click.echo(f"配置文件已保存: {config.config_path}")

    click.echo("系统初始化完成！")


@main.command()
@click.argument("file_path", type=click.Path(exists=True))
@click.pass_context
def parse(ctx, file_path: str):
    """解析单个文件（测试功能）"""
    config = ctx.obj["config"]

    click.echo(f"正在解析文件: {file_path}")

    # 创建解析器
    parser = DocumentParser(config.get_config_dict())

    # 解析文件
    result = parser.parse(file_path)

    if result.success:
        click.echo(f"解析成功！")
        click.echo(f"解析器类型: {result.parser_type}")
        click.echo(f"文档标题: {result.content.title}")
        click.echo(f"文档字数: {result.content.word_count}")
        click.echo(f"文档摘要: {result.summary}")
    else:
        click.echo(f"解析失败: {result.error}", err=True)


@main.command()
@click.argument("source_directory", type=click.Path(exists=True), required=False)
@click.option("--dry-run", is_flag=True, help="仅模拟运行，不实际移动文件")
@click.option("--recursive", "-r", is_flag=True, help="递归处理子目录")
@click.option("--filter-ext", multiple=True, help="只处理指定扩展名的文件")
@click.pass_context
def apply(
    ctx,
    source_directory: Optional[str],
    dry_run: bool,
    recursive: bool,
    filter_ext: tuple,
):
    """执行文档分类整理 (Stage 1 MVP)"""
    config = ctx.obj["config"]

    # 更新配置
    if dry_run:
        config.system.dry_run = True
        click.echo("🔍 模拟运行模式 - 不会实际移动文件")

    # 确定源目录
    if not source_directory:
        source_directory = config.file.source_directory
        if not source_directory:
            click.echo("❌ 未指定源目录，请提供目录路径或在配置中设置", err=True)
            return

    source_path = Path(source_directory)
    if not source_path.exists():
        click.echo(f"❌ 源目录不存在: {source_path}", err=True)
        return

    click.echo(f"📁 处理目录: {source_path}")
    click.echo(f"🎯 目标目录: {config.file.target_directory}")

    try:
        # 创建工作流
        workflow = DocumentClassificationWorkflow(config.get_config_dict())

        # 收集文件
        files_to_process = []

        if recursive:
            pattern = "**/*"
        else:
            pattern = "*"

        # Debug: Print source path and pattern
        click.echo(f"🔍 搜索路径: {source_path} 模式: {pattern}")
        click.echo(f"📋 支持的扩展名: {config.file.supported_extensions}")
        if filter_ext:
            click.echo(f"🎯 过滤扩展名: {filter_ext}")

        for file_path in source_path.glob(pattern):
            if file_path.is_file():
                click.echo(
                    f"📄 发现文件: {file_path.name} (扩展名: {file_path.suffix.lower()})"
                )
                # 检查扩展名过滤
                if filter_ext:
                    # Convert filter_ext to include dots for comparison
                    filter_ext_with_dots = [
                        f".{ext}" if not ext.startswith(".") else ext
                        for ext in filter_ext
                    ]
                    if file_path.suffix.lower() not in filter_ext_with_dots:
                        click.echo(
                            f"   ❌ 被扩展名过滤排除 (过滤列表: {filter_ext_with_dots})"
                        )
                        continue
                else:
                    # 检查是否在支持的扩展名中
                    if file_path.suffix.lower() not in config.file.supported_extensions:
                        click.echo(f"   ❌ 不在支持的扩展名中")
                        continue

                click.echo(f"   ✅ 添加到处理列表")
                files_to_process.append(file_path)

        if not files_to_process:
            click.echo("ℹ️  没有找到需要处理的文件")
            return

        click.echo(f"📋 找到 {len(files_to_process)} 个文件待处理")

        # 处理文件
        results = {
            "total": len(files_to_process),
            "success": 0,
            "failed": 0,
            "needs_review": 0,
            "details": [],
        }

        with click.progressbar(files_to_process, label="处理文件") as files:
            for file_path in files:
                try:
                    result = workflow.process_file(file_path)

                    # 分析结果
                    if result.get("move_success", False):
                        results["success"] += 1
                        status = "✅ 成功"
                    elif result.get("classification", {}).get("needs_review", False):
                        results["needs_review"] += 1
                        status = "⚠️  需要审核"
                    else:
                        results["failed"] += 1
                        status = "❌ 失败"

                    # 记录详细信息
                    details = {
                        "file": str(file_path),
                        "status": status,
                        "category": result.get("classification", {}).get(
                            "primary_category", "未知"
                        ),
                        "confidence": result.get("classification", {}).get(
                            "confidence_score", 0.0
                        ),
                        "new_path": result.get("move_result", {}).get(
                            "primary_target_path", ""
                        ),
                        "error": result.get("error", ""),
                    }
                    results["details"].append(details)

                except Exception as e:
                    results["failed"] += 1
                    click.echo(f"\n❌ 处理失败: {file_path} - {e}")

        # 显示结果汇总
        click.echo("\n" + "=" * 50)
        click.echo("📊 处理结果汇总")
        click.echo("=" * 50)
        click.echo(f"总文件数: {results['total']}")
        click.echo(f"成功处理: {results['success']}")
        click.echo(f"需要审核: {results['needs_review']}")
        click.echo(f"处理失败: {results['failed']}")

        # 显示详细结果
        if ctx.obj.get("verbose", False):
            click.echo("\n📋 详细结果:")
            for detail in results["details"]:
                click.echo(f"{detail['status']} {detail['file']}")
                if detail["category"] != "未知":
                    click.echo(
                        f"    分类: {detail['category']} (置信度: {detail['confidence']:.2f})"
                    )
                if detail["new_path"]:
                    click.echo(f"    新路径: {detail['new_path']}")
                if detail["error"]:
                    click.echo(f"    错误: {detail['error']}")

        # 需要审核的文件
        if results["needs_review"] > 0:
            click.echo(f"\n⚠️  有 {results['needs_review']} 个文件需要人工审核")
            click.echo("💡 使用 'ods review' 命令处理这些文件")

    except Exception as e:
        click.echo(f"❌ 工作流执行失败: {e}", err=True)


@main.command()
@click.pass_context
def info(ctx):
    """显示系统信息"""
    config = ctx.obj["config"]

    click.echo("=== 系统配置信息 ===")
    click.echo(f"配置文件: {config.config_path}")
    click.echo(f"数据库路径: {config.database.path}")
    click.echo(f"临时目录: {config.system.temp_directory}")
    click.echo(f"支持的文件类型: {', '.join(config.file.supported_extensions)}")

    # 解析器信息
    parser = DocumentParser(config.get_config_dict())
    parser_info = parser.get_parser_info()

    click.echo("\n=== 解析器信息 ===")
    click.echo(f"可用解析器: {', '.join(parser_info['available_parsers'])}")
    click.echo(f"支持的扩展名: {', '.join(parser_info['supported_extensions'])}")


@main.command()
@click.argument("source_directory", type=click.Path(exists=True), required=False)
@click.option("--dry-run", is_flag=True, help="仅模拟运行，不实际移动文件")
@click.option("--recursive", "-r", is_flag=True, help="递归处理子目录")
@click.option("--filter-ext", multiple=True, help="只处理指定扩展名的文件")
@click.option("--use-enhanced", is_flag=True, help="使用增强工作流（支持Ollama）")
@click.option("--ollama-only", is_flag=True, help="仅使用Ollama分类器")
@click.pass_context
def apply_enhanced(
    ctx,
    source_directory: Optional[str],
    dry_run: bool,
    recursive: bool,
    filter_ext: tuple,
    use_enhanced: bool,
    ollama_only: bool,
):
    """执行增强文档分类整理（支持Ollama多标签分类）"""
    config = ctx.obj["config"]

    # 更新配置
    if dry_run:
        config.system.dry_run = True
        click.echo("🔍 模拟运行模式 - 不会实际移动文件")

    if ollama_only:
        config["ollama"]["enable_reader"] = True
        click.echo("🤖 仅使用Ollama分类器模式")

    # 确定源目录
    if not source_directory:
        source_directory = config.file.source_directory
        if not source_directory:
            click.echo("❌ 未指定源目录，请提供目录路径或在配置中设置", err=True)
            return

    source_path = Path(source_directory)
    if not source_path.exists():
        click.echo(f"❌ 源目录不存在: {source_path}", err=True)
        return

    click.echo(f"📁 处理目录: {source_path}")
    click.echo(f"🎯 目标目录: {config.file.target_directory}")

    try:
        # 检查是否使用增强工作流
        if use_enhanced or ollama_only:
            click.echo("🚀 使用增强工作流（支持Ollama）")
            from .core.enhanced_workflow import EnhancedWorkflow

            workflow = EnhancedWorkflow(config.get_config_dict())
        else:
            click.echo("⚡ 使用标准工作流")
            workflow = DocumentClassificationWorkflow(config.get_config_dict())

        # 收集文件
        files_to_process = []
        pattern = "**/*" if recursive else "*"

        click.echo(f"🔍 搜索路径: {source_path} 模式: {pattern}")
        click.echo(f"📋 支持的扩展名: {config.file.supported_extensions}")
        if filter_ext:
            click.echo(f"🎯 过滤扩展名: {filter_ext}")

        for file_path in source_path.glob(pattern):
            if file_path.is_file():
                click.echo(
                    f"📄 发现文件: {file_path.name} (扩展名: {file_path.suffix.lower()})"
                )

                # 检查扩展名过滤
                if filter_ext:
                    filter_ext_with_dots = [
                        f".{ext}" if not ext.startswith(".") else ext
                        for ext in filter_ext
                    ]
                    if file_path.suffix.lower() not in filter_ext_with_dots:
                        click.echo(f"   ❌ 被扩展名过滤排除")
                        continue

                # 检查是否在支持的扩展名中
                if file_path.suffix.lower() not in config.file.supported_extensions:
                    click.echo(f"   ❌ 不在支持的扩展名中")
                    continue

                click.echo(f"   ✅ 添加到处理列表")
                files_to_process.append(file_path)

        if not files_to_process:
            click.echo("ℹ️  没有找到需要处理的文件")
            return

        click.echo(f"📋 找到 {len(files_to_process)} 个文件待处理")

        # 处理文件
        results = {
            "total": len(files_to_process),
            "success": 0,
            "failed": 0,
            "needs_review": 0,
            "details": [],
        }

        with click.progressbar(files_to_process, label="处理文件") as files:
            for file_path in files:
                try:
                    result = workflow.process_file(file_path)

                    # 分析结果
                    if result.get("move_success"):
                        results["success"] += 1
                        status = "✅ 成功"
                    elif result.get("needs_review"):
                        results["needs_review"] += 1
                        status = "⚠️  需要审核"
                    else:
                        results["failed"] += 1
                        status = "❌ 失败"

                    # 记录详细信息
                    details = {
                        "file": str(file_path),
                        "status": status,
                        "category": result.get("primary_tag", "未知"),
                        "confidence": result.get("confidence_score", 0.0),
                        "classifier": result.get("classifier_used", "unknown"),
                        "ollama_used": result.get("ollama_content") is not None,
                        "new_path": result.get("final_path", ""),
                        "error": result.get("error", ""),
                        "processing_time": result.get("processing_duration", 0),
                    }
                    results["details"].append(details)

                except Exception as e:
                    results["failed"] += 1
                    click.echo(f"\n❌ 处理失败: {file_path} - {e}")

        # 显示结果汇总
        click.echo("\n" + "=" * 60)
        click.echo("📊 处理结果汇总（增强工作流）")
        click.echo("=" * 60)
        click.echo(f"总文件数: {results['total']}")
        click.echo(f"成功处理: {results['success']}")
        click.echo(f"需要审核: {results['needs_review']}")
        click.echo(f"处理失败: {results['failed']}")

        # 显示详细结果
        if ctx.obj.get("verbose", False):
            click.echo("\n📋 详细结果:")
            for detail in results["details"]:
                click.echo(f"{detail['status']} {detail['file']}")
                if detail["category"] != "未知":
                    click.echo(
                        f"    分类: {detail['category']} (置信度: {detail['confidence']:.2f})"
                    )
                if detail["classifier"] != "unknown":
                    click.echo(f"    分类器: {detail['classifier']}")
                if detail["ollama_used"]:
                    click.echo("    🤖 使用Ollama处理")
                if detail["new_path"]:
                    click.echo(f"    新路径: {detail['new_path']}")
                if detail["processing_time"] > 0:
                    click.echo(f"    处理时间: {detail['processing_time']:.2f}秒")
                if detail["error"]:
                    click.echo(f"    错误: {detail['error']}")

        # 需要审核的文件
        if results["needs_review"] > 0:
            click.echo(f"\n⚠️  有 {results['needs_review']} 个文件需要人工审核")
            click.echo("💡 使用 'ods review' 命令处理这些文件")

        # Ollama状态
        if hasattr(workflow, "get_workflow_summary"):
            summary = workflow.get_workflow_summary()
            click.echo("\n🤖 Ollama状态:")
            click.echo(f"    阅读器启用: {summary.get('ollama_reader_enabled', False)}")
            click.echo(
                f"    分类器启用: {summary.get('ollama_classifier_enabled', False)}"
            )
            click.echo(f"    Ollama可用: {summary.get('ollama_available', False)}")

    except Exception as e:
        click.echo(f"❌ 工作流执行失败: {e}", err=True)


@main.command()
@click.option("--output-dir", "-o", default="./test_files", help="测试文件输出目录")
@click.option("--count", "-c", default=5, help="生成的文件数量")
@click.option("--types", multiple=True, default=["docx", "pdf", "txt"], help="文件类型")
@click.pass_context
def generate_test_files(ctx, output_dir: str, count: int, types: tuple):
    """生成测试文件用于验证系统功能"""
    import os
    from pathlib import Path

    config = ctx.obj["config"]
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    click.echo(f"📁 生成测试文件到: {output_path}")
    click.echo(f"📊 文件数量: {count} 个每种类型")
    click.echo(f"📄 文件类型: {', '.join(types)}")

    # 测试内容模板
    test_contents = [
        {
            "title": "项目计划文档",
            "category": "工作",
            "content": "这是一个项目计划文档，包含了项目目标、时间表和资源分配等重要信息。",
        },
        {
            "title": "会议纪要",
            "category": "工作",
            "content": "本次会议讨论了产品开发进度、质量控制措施以及下阶段工作计划。",
        },
        {
            "title": "个人学习笔记",
            "category": "个人",
            "content": "今天学习了Python编程、机器学习算法和数据结构的相关知识。",
        },
        {
            "title": "财务预算表",
            "category": "财务",
            "content": "本月预算包括人员工资、市场营销费用、办公用品采购等各项支出。",
        },
        {
            "title": "旅行攻略",
            "category": "个人",
            "content": "这次旅行计划去北京，准备参观故宫、天安门和长城等著名景点。",
        },
    ]

    generated_files = []

    try:
        for file_type in types:
            for i in range(count):
                content_idx = i % len(test_contents)
                content = test_contents[content_idx]

                filename = f"{content['category']}_{content['title']}_{i+1}.{file_type}"
                filepath = output_path / filename

                if file_type == "docx":
                    try:
                        from docx import Document as DocxDocument

                        # 生成Word文档
                        doc = DocxDocument()
                        doc.add_heading(content["title"], 0)
                        doc.add_paragraph(content["content"])
                        doc.save(str(filepath))
                    except ImportError:
                        click.echo(f"  ⚠️ 跳过 {filename} - 需要安装 python-docx")
                        continue

                elif file_type == "pdf":
                    try:
                        from reportlab.pdfgen import canvas
                        from reportlab.lib.pagesizes import letter

                        # 生成PDF文档
                        c = canvas.Canvas(str(filepath), pagesize=letter)
                        c.drawString(100, 750, content["title"])
                        c.drawString(100, 700, content["content"])
                        c.save()
                    except ImportError:
                        click.echo(f"  ⚠️ 跳过 {filename} - 需要安装 reportlab")
                        continue

                elif file_type == "txt":
                    # 生成文本文件
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(f"# {content['title']}\n\n{content['content']}\n")

                generated_files.append(str(filepath))
                click.echo(f"  ✅ 生成: {filename}")

        click.echo(f"\n🎉 成功生成 {len(generated_files)} 个测试文件")
        click.echo(f"💡 可以使用以下命令测试文件处理:")
        click.echo(f'   python -m ods apply-enhanced "{output_path}" --use-enhanced')

    except Exception as e:
        click.echo(f"❌ 生成测试文件失败: {e}", err=True)
        raise click.ClickException(str(e))


@main.command()
@click.option("--watch-pid", help="监控进程的PID（用于获取状态）")
@click.pass_context
def watch_status(ctx, watch_pid: Optional[str]):
    """显示监控模式的状态信息"""
    import psutil
    import json

    if watch_pid:
        try:
            pid = int(watch_pid)
            process = psutil.Process(pid)

            click.echo(f"📊 监控进程状态 (PID: {pid}):")
            click.echo(f"   状态: {process.status()}")
            click.echo(f"   CPU使用率: {process.cpu_percent(interval=1):.1f}%")
            click.echo(f"   内存使用: {process.memory_info().rss / 1024 / 1024:.1f} MB")
            click.echo(f"   运行时间: {process.create_time()}")

            # 获取进程命令行
            cmdline = process.cmdline()
            if len(cmdline) > 1:
                click.echo(f"   命令: {' '.join(cmdline[1:])}")

        except (psutil.NoSuchProcess, ValueError):
            click.echo(f"❌ 未找到进程 PID: {watch_pid}")
    else:
        # 显示配置文件中的监控设置
        config_obj = ctx.obj["config"]
        config = config_obj.get_config_dict()
        watcher_config = config.get("watcher", {})

        click.echo("⚙️ 监控配置状态:")
        click.echo(
            f"   启用状态: {'是' if watcher_config.get('enabled', True) else '否'}"
        )
        click.echo(f"   检测间隔: {watcher_config.get('check_interval', 5)}秒")
        click.echo(f"   递归监控: {'是' if watcher_config.get('recursive', True) else '否'}")
        click.echo(f"   队列大小: {watcher_config.get('max_queue_size', 100)}")
        click.echo(f"   去抖时间: {watcher_config.get('debounce_time', 2)}秒")

        strategy = watcher_config.get("strategy", {})
        click.echo(f"   并发线程: {strategy.get('workers', 2)}")
        click.echo(f"   批量大小: {strategy.get('batch_size', 5)}")
        click.echo(f"   批量处理: {'是' if strategy.get('batch_process', True) else '否'}")

        filters = watcher_config.get("file_filters", {})
        if filters.get("extensions"):
            click.echo(f"   文件过滤: {', '.join(filters['extensions'])}")
        else:
            click.echo("   文件过滤: 所有支持的格式")

        click.echo(f"   大小限制: {filters.get('min_size', 0)} - {filters.get('max_size', 104857600)} bytes")


@main.command()
@click.argument("directory", type=click.Path(exists=True))
@click.option("--recursive", "-r", is_flag=True, help="递归检查子目录")
@click.pass_context
def validate_files(ctx, directory: str, recursive: bool):
    """验证目录中文件的完整性"""
    from pathlib import Path
    from ..parsers.office_parser import OfficeParser

    config = ctx.obj["config"]
    parser = OfficeParser(config.get("parsers", {}))
    directory_path = Path(directory)

    click.echo(f"🔍 验证目录: {directory_path}")
    click.echo(f"🔄 递归检查: {'是' if recursive else '否'}")

    pattern = "**/*" if recursive else "*"
    files = list(directory_path.glob(pattern))

    # 过滤出支持的文件类型
    supported_extensions = {
        ".docx",
        ".doc",
        ".pptx",
        ".ppt",
        ".xls",
        ".xlsx",
        ".pdf",
        ".txt",
        ".md",
    }
    target_files = [
        f for f in files if f.is_file() and f.suffix.lower() in supported_extensions
    ]

    if not target_files:
        click.echo("❌ 未找到支持的文件类型")
        return

    click.echo(f"📁 找到 {len(target_files)} 个文件待验证")

    results = {"good": 0, "corrupted": 0, "empty": 0, "error": 0}

    for filepath in target_files:
        try:
            report = parser.check_file_integrity(filepath)

            if report["integrity_status"] == "good":
                results["good"] += 1
                status = "✅"
            elif report["integrity_status"] == "corrupted":
                results["corrupted"] += 1
                status = "❌"
            elif report["integrity_status"] == "empty":
                results["empty"] += 1
                status = "📭"
            else:
                results["error"] += 1
                status = "⚠️"

            click.echo(f"  {status} {filepath.name}")
            if report["issues"]:
                click.echo(f"      问题: {', '.join(report['issues'])}")
            if report["recommendations"]:
                click.echo(f"      建议: {', '.join(report['recommendations'])}")

        except Exception as e:
            results["error"] += 1
            click.echo(f"  ⚠️ {filepath.name} - 验证失败: {e}")

    # 显示汇总
    click.echo(f"\n📊 验证结果汇总:")
    click.echo(f"  ✅ 正常文件: {results['good']} 个")
    click.echo(f"  ❌ 损坏文件: {results['corrupted']} 个")
    click.echo(f"  📭 空文件: {results['empty']} 个")
    click.echo(f"  ⚠️ 验证失败: {results['error']} 个")

    total_files = sum(results.values())
    healthy_percentage = (results["good"] / total_files * 100) if total_files > 0 else 0

    if healthy_percentage > 80:
        click.echo(f"🎉 文件健康度: {healthy_percentage:.1f}% - 系统运行良好")
    elif healthy_percentage > 50:
        click.echo(f"⚠️ 文件健康度: {healthy_percentage:.1f}% - 部分文件存在问题")
    else:
        click.echo(f"❌ 文件健康度: {healthy_percentage:.1f}% - 大多数文件存在问题")
        click.echo(
            f"💡 建议: 检查文件来源或使用 'generate-test-files' 命令创建测试文件"
        )


@main.command()
@click.pass_context
def check_ollama(ctx):
    """检查Ollama服务状态和可用模型"""
    import requests

    config_obj = ctx.obj["config"]
    config = config_obj.get_config_dict()
    ollama_config = config.get("ollama", {})
    base_url = ollama_config.get("base_url", "http://localhost:11434")

    click.echo("🔍 检查Ollama服务状态...")
    click.echo(f"📍 服务地址: {base_url}")

    try:
        # 检查服务是否运行
        response = requests.get(f"{base_url}/api/tags", timeout=10)

        if response.status_code == 200:
            click.echo("✅ Ollama服务运行正常")

            # 获取可用模型
            data = response.json()
            models = data.get("models", [])

            if models:
                click.echo(f"\n📋 可用模型 ({len(models)} 个):")
                for model in models:
                    name = model.get("name", "未知")
                    size = model.get("size", 0)
                    size_gb = size / (1024**3) if size else 0
                    click.echo(f"  {name} ({size_gb:.1f} GB)")
            else:
                click.echo("\n⚠️  没有找到可用模型")
                click.echo("💡 请运行以下命令安装模型:")
                click.echo("   ollama pull qwen2.5:3b")
                click.echo("   ollama pull qwen2.5:7b")

        else:
            click.echo(f"❌ Ollama服务响应异常: {response.status_code}")

    except requests.exceptions.ConnectionError:
        click.echo("❌ 无法连接到Ollama服务")
        click.echo("💡 请确保Ollama服务正在运行:")
        click.echo("   1. 安装Ollama: https://ollama.com/download")
        click.echo("   2. 启动服务: ollama serve")
        click.echo("   3. 安装模型: ollama pull qwen2.5:3b")

    except Exception as e:
        click.echo(f"❌ 检查失败: {e}")

    # 显示配置信息
    click.echo(f"\n⚙️  当前配置:")
    click.echo(f"   模型: {ollama_config.get('model', '未设置')}")
    click.echo(f"   阅读模型: {ollama_config.get('reader_model', '未设置')}")
    click.echo(f"   分类模型: {ollama_config.get('classifier_model', '未设置')}")
    click.echo(f"   超时时间: {ollama_config.get('timeout', 300)} 秒")
    click.echo(f"   最大重试: {ollama_config.get('max_retries', 3)} 次")


@main.command()
@click.argument("source_directory", type=click.Path(exists=True), required=False)
@click.option("--recursive", "-r", is_flag=True, help="递归监控子目录")
@click.option("--use-enhanced", is_flag=True, help="使用增强工作流（支持Ollama）")
@click.option("--ollama-only", is_flag=True, help="仅使用Ollama分类器")
@click.option("--interval", "-i", default=5, help="文件变化检测间隔（秒）")
@click.option("--quiet", "-q", is_flag=True, help="静默模式，不显示详细信息")
@click.option("--filter-ext", multiple=True, help="只监控指定扩展名的文件")
@click.option("--workers", "-w", default=2, help="并发处理的工作线程数")
@click.option("--batch-size", "-b", default=5, help="批量处理文件数量")
@click.pass_context
def watch(
    ctx,
    source_directory: Optional[str],
    recursive: bool,
    use_enhanced: bool,
    ollama_only: bool,
    interval: int,
    quiet: bool,
    filter_ext: tuple,
    workers: int,
    batch_size: int,
):
    """启动后台监控模式，自动处理新增或修改的文件"""
    import os
    import signal
    import time
    import threading
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from pathlib import Path
    from .core.watcher import DirectoryWatcher

    config_obj = ctx.obj["config"]
    config_dict = config_obj.get_config_dict()

    # 获取监控配置
    watcher_config = config_dict.get("watcher", {})

    # 使用配置中的默认值（如果用户没有明确指定）
    # 对于有默认值的参数，我们总是尝试从配置中获取更好的值
    interval = watcher_config.get("check_interval", interval)
    if not recursive:
        recursive = watcher_config.get("recursive", True)

    # 使用配置中的并发设置
    strategy = watcher_config.get("strategy", {})
    workers = strategy.get("workers", workers)
    batch_size = strategy.get("batch_size", batch_size)

    # 确定源目录
    if not source_directory:
        source_directory = config_obj.file.source_directory
        if not source_directory:
            click.echo("❌ 未指定源目录，请提供目录路径或在配置中设置", err=True)
            return

    source_path = Path(source_directory)
    if not source_path.exists():
        click.echo(f"❌ 源目录不存在: {source_path}", err=True)
        return

    click.echo(f"📁 监控目录: {source_path}")
    click.echo(f"🔄 递归监控: {'是' if recursive else '否'}")
    click.echo(f"⏰ 检测间隔: {interval}秒")

    # 显示文件过滤信息
    if filter_ext:
        click.echo(f"📄 文件过滤: {', '.join(filter_ext)}")
    elif watcher_config.get("file_filters", {}).get("extensions"):
        click.echo(f"📄 文件过滤: {', '.join(watcher_config['file_filters']['extensions'])}")

    click.echo(f"🚀 工作流: {'增强模式（Ollama）' if use_enhanced or ollama_only else '标准模式'}")

    # 显示监控配置信息
    if not quiet:
        click.echo(f"\n⚙️ 监控配置:")
        click.echo(f"   队列大小: {watcher_config.get('max_queue_size', 100)}")
        click.echo(f"   去抖时间: {watcher_config.get('debounce_time', 2)}秒")
        click.echo(f"   批量处理: {'是' if strategy.get('batch_process', True) else '否'}")
        click.echo(f"   并发线程: {workers}")
        click.echo(f"   批处理大小: {batch_size}")
    click.echo("")

    # 初始化工作流
    try:
        if use_enhanced or ollama_only:
            from .core.enhanced_workflow import EnhancedWorkflow
            workflow = EnhancedWorkflow(config_dict)
        else:
            from .core.workflow import DocumentClassificationWorkflow
            workflow = DocumentClassificationWorkflow(config_dict)
    except Exception as e:
        click.echo(f"❌ 工作流初始化失败: {e}", err=True)
        return

    # 待处理文件队列和状态跟踪
    pending_files = set()
    processed_files = set()
    processing_lock = threading.Lock()  # 线程安全锁
    processing_stats = {
        "processed": 0,
        "successful": 0,
        "failed": 0,
        "skipped": 0
    }

    def on_file_change(file_path: str):
        """文件变化回调函数"""
        file_path_obj = Path(file_path)

        # 检查文件是否存在且是文件
        if not file_path_obj.exists() or file_path_obj.is_dir():
            return

        # 检查文件扩展名过滤
        supported_extensions = {'.docx', '.doc', '.pptx', '.ppt', '.xls', '.xlsx',
                               '.pdf', '.txt', '.md', '.jpg', '.jpeg', '.png'}

        # 使用命令行参数或配置文件中的扩展名过滤
        if filter_ext:
            target_extensions = [ext.lower() for ext in filter_ext]
        else:
            file_filters = watcher_config.get("file_filters", {})
            target_extensions = [ext.lower() for ext in file_filters.get("extensions", [])]
            if not target_extensions:
                target_extensions = list(supported_extensions)

        if file_path_obj.suffix.lower() not in target_extensions:
            return

        # 检查文件大小过滤
        file_filters = watcher_config.get("file_filters", {})
        min_size = file_filters.get("min_size", 0)
        max_size = file_filters.get("max_size", 104857600)  # 100MB

        try:
            file_size = file_path_obj.stat().st_size
            if file_size < min_size or file_size > max_size:
                if not quiet:
                    click.echo(f"📄 文件大小不符合要求，跳过: {file_path_obj.name} ({file_size} bytes)")
                return
        except OSError:
            # 如果无法获取文件大小，跳过
            return

        # 检查是否是已处理文件（避免重复处理）
        if str(file_path_obj) in processed_files:
            return

        # 检查队列大小限制
        max_queue_size = watcher_config.get("max_queue_size", 100)
        if len(pending_files) >= max_queue_size:
            if not quiet:
                click.echo(f"⚠️ 队列已满，暂时跳过文件: {file_path_obj.name}")
            return

        # 添加到待处理队列
        with processing_lock:
            pending_files.add(str(file_path_obj))
        if not quiet:
            click.echo(f"📄 检测到文件变化: {file_path_obj.name}")

    def process_file_concurrent(file_path: str) -> dict:
        """并发处理单个文件的函数"""
        try:
            file_path_obj = Path(file_path)

            # 检查文件是否仍然存在且不是目录
            if not file_path_obj.exists() or file_path_obj.is_dir():
                with processing_lock:
                    processing_stats["skipped"] += 1
                return {"status": "skipped", "reason": "文件不存在或已删除"}

            # 等待文件稳定（避免处理正在写入的文件）
            debounce_time = watcher_config.get("debounce_time", 2)
            time.sleep(debounce_time)

            # 处理文件
            result = workflow.process_file(file_path_obj)

            # 更新统计信息
            with processing_lock:
                processing_stats["processed"] += 1
                if result.get("status") == "completed":
                    processing_stats["successful"] += 1
                else:
                    processing_stats["failed"] += 1

            return result

        except Exception as e:
            with processing_lock:
                processing_stats["failed"] += 1
                processing_stats["processed"] += 1
            return {"status": "error", "error": str(e)}

    # 创建目录监听器
    watcher = DirectoryWatcher(str(source_path), on_file_change, recursive=recursive)

    # 处理信号以优雅关闭
    shutdown_requested = False
    shutdown_reason = None

    def signal_handler(signum, frame):
        nonlocal shutdown_requested, shutdown_reason
        if signum == signal.SIGINT:
            shutdown_reason = "用户中断 (Ctrl+C)"
        elif signum == signal.SIGTERM:
            shutdown_reason = "系统终止信号"
        else:
            shutdown_reason = f"信号 {signum}"

        click.echo(f"\n\n🛑 收到关闭信号 ({shutdown_reason})，正在停止监控...")
        shutdown_requested = True

    signal.signal(signal.SIGINT, signal_handler)
    signal.signal(signal.SIGTERM, signal_handler)

    # 处理其他常见信号（仅在支持的平台上）
    def handle_other_signals(signum, frame):
        nonlocal shutdown_requested, shutdown_reason
        shutdown_reason = f"系统信号 {signum}"
        click.echo(f"\n\n🛑 收到系统信号，正在安全关闭...")
        shutdown_requested = True

    # 仅在支持的平台上注册信号处理器
    import platform

    if platform.system() != "Windows":
        # Unix/Linux 特有的信号
        if hasattr(signal, "SIGHUP"):
            signal.signal(signal.SIGHUP, handle_other_signals)  # 终端关闭
        if hasattr(signal, "SIGUSR1"):
            signal.signal(signal.SIGUSR1, handle_other_signals)  # 用户信号1

    try:
        # 启动监听器
        watcher.start()
        click.echo("✅ 文件监控已启动")
        click.echo(f"📍 进程PID: {os.getpid()}")
        click.echo("💡 使用以下命令查看状态:")
        click.echo(f"   python -m ods watch-status --watch-pid {os.getpid()}")
        click.echo("💡 按 Ctrl+C 停止监控")
        click.echo("-" * 50)

        # 主监控循环
        with ThreadPoolExecutor(max_workers=workers) as executor:
            while not shutdown_requested:
                # 获取待处理文件
                with processing_lock:
                    files_to_process = list(pending_files)[:batch_size]  # 限制批量大小
                    for file_path in files_to_process:
                        pending_files.discard(file_path)

                if files_to_process:
                    if not quiet:
                        click.echo(f"🔄 批量处理 {len(files_to_process)} 个文件...")

                    # 提交并发任务
                    future_to_file = {
                        executor.submit(process_file_concurrent, file_path): file_path
                        for file_path in files_to_process
                    }

                    # 等待任务完成并处理结果
                    for future in as_completed(future_to_file):
                        file_path = future_to_file[future]
                        try:
                            result = future.result()

                            # 记录处理结果
                            with processing_lock:
                                processed_files.add(file_path)

                            if not quiet:
                                file_name = Path(file_path).name
                                if result.get("status") == "completed":
                                    click.echo(f"✅ 处理完成: {file_name}")
                                elif result.get("status") == "skipped":
                                    click.echo(f"⏭️ 跳过处理: {file_name}")
                                else:
                                    error_msg = result.get('error', '未知错误')
                                    click.echo(f"❌ 处理失败: {file_name} - {error_msg}")

                        except Exception as e:
                            if not quiet:
                                click.echo(f"❌ 任务异常: {Path(file_path).name} - {e}")

                    # 显示状态
                    with processing_lock:
                        if not quiet and processing_stats["processed"] > 0 and processing_stats["processed"] % (batch_size * 2) == 0:
                            click.echo(f"📊 处理统计: {processing_stats['processed']} 已处理, "
                                     f"{processing_stats['successful']} 成功, "
                                     f"{processing_stats['failed']} 失败, "
                                     f"{processing_stats['skipped']} 跳过")

                time.sleep(interval)

    except KeyboardInterrupt:
        if not shutdown_reason:
            shutdown_reason = "用户中断 (Ctrl+C)"
    except Exception as e:
        shutdown_reason = f"运行时错误: {e}"
        click.echo(f"\n❌ 监控过程出错: {e}")
    finally:
        # 优雅关闭流程
        click.echo("\n🔄 开始清理资源...")

        # 1. 停止接收新文件
        click.echo("   停止文件监听...")
        watcher.stop()

        # 2. 完成当前处理的任务
        click.echo("   等待正在处理的文件完成...")
        # 线程池会在with语句结束后自动清理

        # 3. 等待监听器完全停止
        click.echo("   等待监听器关闭...")
        watcher.join(timeout=5)  # 最多等待5秒

        # 4. 显示最终统计信息
        with processing_lock:
            if processing_stats["processed"] > 0:
                click.echo(f"\n📊 最终处理统计:")
                click.echo(f"   总文件数: {processing_stats['processed']}")
                click.echo(f"   成功处理: {processing_stats['successful']}")
                click.echo(f"   处理失败: {processing_stats['failed']}")
                click.echo(f"   跳过处理: {processing_stats['skipped']}")

                success_rate = (processing_stats['successful'] / processing_stats['processed'] * 100) if processing_stats['processed'] > 0 else 0
                click.echo(f"   成功率: {success_rate:.1f}%")

                # 显示关闭原因
                if shutdown_reason:
                    click.echo(f"   关闭原因: {shutdown_reason}")

        click.echo("✅ 监控已完全停止")
        click.echo(f"💡 提示: 如需重新启动，使用以下命令:")
        click.echo(f"   python -m ods watch \"{source_path}\" {'--recursive' if recursive else ''} {'--use-enhanced' if use_enhanced or ollama_only else ''}")


@main.command()
@click.option("--max-files", "-n", default=10, help="最大审核文件数")
@click.option("--user-id", "-u", help="用户ID（可选）")
@click.option("--batch", "-b", is_flag=True, help="启用批量审核模式")
@click.pass_context
def review(ctx, max_files: int, user_id: str, batch: bool):
    """启动交互式文件审核界面"""
    config_obj = ctx.obj["config"]
    config = config_obj.get_config_dict()

    try:
        # 导入审核模块
        from .review.interactive_reviewer import InteractiveReviewer

        # 创建审核界面
        reviewer = InteractiveReviewer(config)

        # 检查是否有待审核文件
        pending_count = reviewer.get_pending_reviews_count()
        if pending_count == 0:
            click.echo("✅ 没有找到需要审核的文件！")
            click.echo(
                "💡 提示: 运行 'ods apply' 进行文件分类，系统会自动标记需要审核的文件"
            )
            return

        click.echo(f"📋 发现 {pending_count} 个待审核文件")

        if batch:
            click.echo("🔄 启用批量审核模式")
            click.echo("💡 提示: 批量模式可以对多个文件应用相同的操作，提高效率")

        # 开始审核会话
        session_id = reviewer.start_review_session(user_id)

        # 运行交互式审核
        reviewer.run_interactive_review(session_id, max_files, batch_mode=batch)

    except ImportError as e:
        click.echo(f"❌ 无法加载审核模块: {e}", err=True)
        click.echo("💡 请确保所有依赖都已正确安装", err=True)
    except Exception as e:
        click.echo(f"❌ 审核过程出错: {e}", err=True)
        click.echo("💡 请检查日志文件获取详细信息", err=True)


@main.command()
@click.option("--session-id", "-s", help="审核会话ID")
@click.option("--detailed", "-d", is_flag=True, help="显示详细统计信息")
@click.pass_context
def review_stats(ctx, session_id: str, detailed: bool):
    """查看审核统计信息"""
    config_obj = ctx.obj["config"]
    config = config_obj.get_config_dict()

    try:
        # 导入审核管理器
        from .review.review_manager import ReviewManager

        # 创建审核管理器
        review_manager = ReviewManager(config)

        # 获取统计信息
        stats = review_manager.get_review_statistics(session_id)

        if not stats:
            click.echo("❌ 未找到审核统计信息")
            return

        click.echo("\n📊 审核统计信息")
        click.echo("=" * 50)

        if session_id:
            # 单个会话统计
            session_info = stats.get("session", {})
            records_info = stats.get("records", {})

            click.echo(f"🎯 会话ID: {session_id}")
            click.echo(f"👤 用户: {session_info.get('user_id', '未指定')}")
            click.echo(f"📅 开始时间: {session_info.get('start_time', '未知')}")
            click.echo(f"📊 总文件数: {session_info.get('total_files', 0)}")
            click.echo(f"✅ 已审核: {session_info.get('reviewed_files', 0)}")
            click.echo(f"📈 完成率: {stats.get('completion_rate', 0):.1f}%")
            click.echo(f"📋 审核记录: {records_info.get('total_reviews', 0)}")
            click.echo(f"  ✅ 批准: {records_info.get('approved', 0)}")
            click.echo(f"  ✏️  修改: {records_info.get('corrected', 0)}")
            click.echo(f"  🚫 拒绝: {records_info.get('rejected', 0)}")

            if detailed:
                avg_time = records_info.get("avg_processing_time", 0)
                if avg_time:
                    click.echo(f"  ⏱️  平均处理时间: {avg_time:.2f} 秒")
        else:
            # 全局统计
            click.echo(f"📂 待审核文件: {stats.get('pending_reviews', 0)}")
            click.echo(f"📊 审核会话总数: {stats.get('total_sessions', 0)}")

            review_actions = stats.get("review_actions", {})
            if review_actions:
                click.echo(f"📋 审核记录总数: {review_actions.get('total', 0)}")
                click.echo(f"  ✅ 批准: {review_actions.get('approved', 0)}")
                click.echo(f"  ✏️  修改: {review_actions.get('corrected', 0)}")
                click.echo(f"  🚫 拒绝: {review_actions.get('rejected', 0)}")

    except ImportError as e:
        click.echo(f"❌ 无法加载审核模块: {e}", err=True)
    except Exception as e:
        click.echo(f"❌ 获取统计信息出错: {e}", err=True)


if __name__ == "__main__":
    main()
